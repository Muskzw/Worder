import uuid
import os
from flask import Flask, request, send_file, render_template_string, redirect, url_for, flash
from pdf2docx import Converter
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
from docx import Document
from werkzeug.utils import secure_filename
import camelot
from docx.shared import Inches

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10 MB limit
app.secret_key = 'your_secret_key'  # Needed for flashing messages

HTML_FORM = '''
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>File to Word Converter</title>
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <style>
    body {
      background: linear-gradient(135deg, #e0e7ff 0%, #f4f4f4 100%);
      color: #222;
      font-family: 'Segoe UI', Arial, sans-serif;
      margin: 0;
      min-height: 100vh;
      transition: background 0.3s, color 0.3s;
      display: flex;
      flex-direction: column;
      align-items: center;
      justify-content: center;
    }
    .dark-mode {
      background: linear-gradient(135deg, #232323 0%, #181818 100%);
      color: #eee;
    }
    .toggle-btn {
      position: fixed;
      top: 24px;
      right: 24px;
      padding: 10px 22px;
      background: #222;
      color: #fff;
      border: none;
      border-radius: 22px;
      font-size: 1rem;
      cursor: pointer;
      box-shadow: 0 2px 8px rgba(0,0,0,0.08);
      transition: background 0.3s, color 0.3s;
      z-index: 10;
    }
    .dark-mode .toggle-btn {
      background: #eee;
      color: #222;
    }
    .container {
      max-width: 420px;
      margin: 80px auto 0 auto;
      padding: 36px 32px 28px 32px;
      background: rgba(255,255,255,0.25);
      border-radius: 20px;
      box-shadow: 0 8px 32px 0 rgba(31,38,135,0.18);
      text-align: center;
      transition: background 0.3s, box-shadow 0.3s;
      backdrop-filter: blur(12px);
      -webkit-backdrop-filter: blur(12px);
      border: 1px solid rgba(255,255,255,0.18);
    }
    .dark-mode .container {
      background: rgba(35,35,35,0.45);
      box-shadow: 0 8px 32px 0 rgba(0,0,0,0.35);
      border: 1px solid rgba(255,255,255,0.08);
    }
    h2 {
      margin-top: 0;
      font-weight: 600;
      font-size: 1.6rem;
      letter-spacing: 0.5px;
    }
    p {
      color: #555;
      margin-bottom: 28px;
      font-size: 1.05rem;
    }
    .dark-mode p {
      color: #bbb;
    }
    input[type="file"] {
      margin: 18px 0 18px 0;
      font-size: 1rem;
    }
    label, select {
      font-size: 1rem;
      margin-bottom: 12px;
    }
    button[type="submit"] {
      margin-top: 18px;
      padding: 10px 32px;
      background: #0074d9;
      color: #fff;
      border: none;
      border-radius: 22px;
      font-size: 1.1rem;
      font-weight: 500;
      cursor: pointer;
      box-shadow: 0 2px 8px rgba(0,0,0,0.08);
      transition: background 0.2s;
    }
    button[type="submit"]:hover {
      background: #005fa3;
    }
    .dark-mode button[type="submit"] {
      background: #3399ff;
      color: #181818;
    }
    .dark-mode button[type="submit"]:hover {
      background: #0074d9;
      color: #fff;
    }
    ul {
      padding-left: 0;
      list-style: none;
      margin: 0 0 18px 0;
    }
    li {
      color: #e74c3c;
      font-size: 1rem;
      margin-bottom: 6px;
    }
    a {
      color: #0074d9;
      text-decoration: none;
      font-weight: 500;
      transition: color 0.2s;
    }
    .dark-mode a {
      color: #66b3ff;
    }
    .success-block {
      display: flex;
      flex-direction: column;
      align-items: center;
      justify-content: center;
      margin-top: 36px;
      margin-bottom: 12px;
      padding: 24px 18px;
      background: rgba(255,255,255,0.45);
      border-radius: 16px;
      box-shadow: 0 2px 12px rgba(0,0,0,0.07);
      backdrop-filter: blur(8px);
      -webkit-backdrop-filter: blur(8px);
      border: 1px solid rgba(255,255,255,0.18);
    }
    .dark-mode .success-block {
      background: rgba(35,35,35,0.55);
      border: 1px solid rgba(255,255,255,0.08);
    }
    #progress-bar {
      display: none;
      margin: 20px 0;
    }
    #progress-bar div {
      width: 100%;
      background: #e0e0e0;
      border-radius: 8px;
    }
    #bar {
      width: 0%;
      height: 18px;
      background: #0074d9;
      border-radius: 8px;
      transition: width 0.4s;
    }
    #progress-text {
      text-align: center;
      margin-top: 6px;
      color: #555;
    }
    @media (max-width: 600px) {
      .container {
        padding: 18px 6vw 18px 6vw;
        margin-top: 32px;
      }
      .toggle-btn {
        top: 12px;
        right: 12px;
        padding: 8px 14px;
        font-size: 0.95rem;
      }
      .success-block {
        padding: 14px 4vw;
      }
    }
  </style>
</head>
<body>
  <button class="toggle-btn" onclick="toggleDarkMode()">🌙 Toggle Dark Mode</button>
  <div class="container">
    <h2>File to Word Converter</h2>
    <p>
      Upload a PDF, image, or text file and convert it to a Word document.<br>
      <span style="font-size:0.97em;">Choose OCR language for images.</span>
    </p>
    {% with messages = get_flashed_messages() %}
      {% if messages %}
        <ul>
        {% for message in messages %}
          <li>{{ message }}</li>
        {% endfor %}
        </ul>
      {% endif %}
    {% endwith %}
    <form method="post" action="/convert" enctype="multipart/form-data">
      <input type="file" name="file" required><br>
      <label for="lang">OCR Language (for images):</label>
      <select name="lang" id="lang">
        <option value="eng">English</option>
        <option value="spa">Spanish</option>
      </select><br>
      <label>
        <input type="checkbox" name="extract_tables" value="yes">
        Extract tables from PDF (experimental)
      </label><br>
      <button type="submit">Convert</button>
    </form>
    <div id="progress-bar">
      <div style="width:100%;background:#e0e0e0;border-radius:8px;">
        <div id="bar" style="width:0%;height:18px;background:#0074d9;border-radius:8px;transition:width 0.4s;"></div>
      </div>
      <div style="text-align:center;margin-top:6px;color:#555;" id="progress-text">Converting...</div>
    </div>
  </div>
  <script>
    function toggleDarkMode() {
      document.body.classList.toggle('dark-mode');
      localStorage.setItem('darkMode', document.body.classList.contains('dark-mode'));
    }
    // On load, set mode from localStorage
    if (localStorage.getItem('darkMode') === 'true') {
      document.body.classList.add('dark-mode');
    }
  </script>
</body>
</html>
'''

@app.route('/')
def home():
    download_filename = request.args.get('download')
    return render_template_string(
        HTML_FORM + '''
        {% if download_filename %}
          <div class="success-block">
            <h3 style="margin-bottom:12px;">✅ Conversion successful!</h3>
            <a href="{{ url_for('download', filename=download_filename) }}" style="font-size:1.15rem;">⬇️ Download your Word file</a>
          </div>
        {% endif %}
        ''',
        download_filename=download_filename
    )

@app.route('/convert', methods=['POST'])
def convert():
    uploaded_file = request.files['file']
    if not uploaded_file:
        flash("No file uploaded")
        return redirect(url_for('home'))

    filename = secure_filename(uploaded_file.filename)
    if not filename.lower().endswith(('.pdf', '.png', '.jpg', '.jpeg', '.txt')):
        flash("Unsupported file type")
        return redirect(url_for('home'))

    unique_id = str(uuid.uuid4())
    filepath = os.path.join(UPLOAD_FOLDER, unique_id + "_" + filename)
    uploaded_file.save(filepath)

    output_filename = unique_id + "_converted.docx"
    output_path = os.path.join(UPLOAD_FOLDER, output_filename)

    try:
        if filename.lower().endswith('.pdf'):
            try:
                cv = Converter(filepath)
                cv.convert(output_path, start=0, end=None)
                cv.close()
            except Exception:
                lang = request.form.get('lang', 'eng')
                images = convert_from_path(filepath)
                doc = Document()
                for img in images:
                    text = pytesseract.image_to_string(img, lang=lang)
                    doc.add_paragraph(text)
                doc.save(output_path)
        elif filename.lower().endswith(('.png', '.jpg', '.jpeg')):
            lang = request.form.get('lang', 'eng')
            text = pytesseract.image_to_string(Image.open(filepath), lang=lang)
            doc = Document()
            doc.add_paragraph(text)
            doc.save(output_path)
        elif filename.lower().endswith('.txt'):
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            doc = Document()
            doc.add_paragraph(content)
            doc.save(output_path)
    except Exception as e:
        flash(f"Conversion failed: {e}")
        if os.path.exists(filepath):
            os.remove(filepath)
        return redirect(url_for('home'))

    extract_tables = request.form.get('extract_tables') == 'yes'
    if extract_tables and filename.lower().endswith('.pdf'):
        try:
            tables = camelot.read_pdf(filepath, pages='all')
            if tables:
                doc = Document(output_path)
                doc.add_page_break()
                doc.add_heading('Extracted Tables', level=2)
                for i, table in enumerate(tables):
                    doc.add_paragraph(f"Table {i+1}:")
                    data = table.df.values.tolist()
                    rows, cols = len(data), len(data[0])
                    table_docx = doc.add_table(rows=rows, cols=cols)
                    for r in range(rows):
                        for c in range(cols):
                            table_docx.cell(r, c).text = str(data[r][c])
                    doc.add_paragraph("")
                doc.save(output_path)
        except Exception as e:
            flash(f"Table extraction failed: {e}")

    if os.path.exists(filepath):
        os.remove(filepath)

    # Redirect to home with download link
    return redirect(url_for('home', download=output_filename))

@app.route('/download/<filename>')
def download(filename):
    output_path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(output_path):
        return "File not found", 404
    response = send_file(output_path, as_attachment=True)
    # Remove the file after sending
    try:
        os.remove(output_path)
    except Exception:
        pass
    return response

if __name__ == '__main__':
    app.run(debug=True)
